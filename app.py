# Multimodal RAG (Text · Image · Audio) — NO external FFmpeg
# PyAV (av) decodes audio to 16k mono numpy for Whisper. CLIP (text+image) shared space, FAISS retrieval,
# grounded answers with citations. Robust YouTube pipeline with format fallback & validation.

import os, io, uuid, json, tempfile
from typing import List, Dict, Any, Tuple

import streamlit as st
import numpy as np
from PIL import Image

# Embeddings / index
from sentence_transformers import SentenceTransformer
import faiss

# Loaders
from langchain_community.document_loaders import PyPDFLoader, TextLoader
try:
    from langchain_community.document_loaders import Docx2txtLoader
    HAS_DOCX2TXT = True
except Exception:
    HAS_DOCX2TXT = False
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.llms import Ollama

# Web & YouTube
import requests
from bs4 import BeautifulSoup
try:
    import yt_dlp
    HAS_YTDLP = True
except Exception:
    HAS_YTDLP = False

# Audio (FFmpeg-free)
import av                      # PyAV brings ffmpeg libs in wheel
import soundfile as sf         # write wavs for playback
try:
    import whisper
    HAS_WHISPER = True
except Exception:
    HAS_WHISPER = False

st.set_page_config(page_title="SIH Multimodal RAG — no external FFmpeg", layout="wide")
st.title("🧭 SIH PS-25231 · Multimodal RAG (Text · Image · Audio) — no external FFmpeg")

# ---------------- Models ----------------
@st.cache_resource(show_spinner=False)
def load_clip_model():
    return SentenceTransformer("clip-ViT-B-32")

@st.cache_resource(show_spinner=False)
def load_whisper_model():
    if not HAS_WHISPER:
        raise RuntimeError("Whisper not installed. Run: pip install whisper openai-whisper")
    return whisper.load_model("base")

def embed_text(clip, text: str) -> np.ndarray:
    return clip.encode(text, normalize_embeddings=True).astype("float32")

def embed_image(clip, img: Image.Image) -> np.ndarray:
    return clip.encode(img.convert("RGB"), normalize_embeddings=True).astype("float32")

# ---------------- Audio helpers (PyAV) ----------------
TARGET_SR = 16000

def decode_audio_to_16k_mono(path: str, max_seconds: float | None = None) -> np.ndarray:
    """
    Decode any common audio container/codec using PyAV; resample to 16 kHz mono; return float32 numpy [-1,1].
    Handles resampler returning a single frame OR a list of frames. Early-stops at max_seconds.
    """
    container = av.open(path)
    if not container.streams.audio:
        container.close()
        return np.zeros(0, dtype=np.float32)

    astream = container.streams.audio[0]
    resampler = av.audio.resampler.AudioResampler(format="s16", layout="mono", rate=TARGET_SR)

    chunks = []
    total_samples = 0
    limit_samples = int(max_seconds * TARGET_SR) if max_seconds else None

    for frame in container.decode(astream):
        outs = resampler.resample(frame)
        if not isinstance(outs, list):
            outs = [outs]
        for out in outs:
            if out is None:
                continue
            pcm = out.to_ndarray()            # (1, N) int16
            chunks.append(pcm)
            total_samples += pcm.shape[1]
            if limit_samples is not None and total_samples >= limit_samples:
                break
        if limit_samples is not None and total_samples >= limit_samples:
            break

    container.close()

    if not chunks:
        return np.zeros(0, dtype=np.float32)

    cat = np.concatenate(chunks, axis=1)
    if limit_samples is not None:
        cat = cat[:, :limit_samples]
    pcm16 = cat.ravel().astype(np.int16)
    audio = (pcm16.astype(np.float32) / 32768.0).copy()
    return audio

def write_wav(path: str, audio: np.ndarray, sr: int = TARGET_SR):
    sf.write(path, audio, sr, subtype="PCM_16")

# ---------------- DOCX fallback ----------------
def load_docx_via_python_docx(path: str) -> List[Document]:
    try:
        import docx
    except Exception as e:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx") from e
    doc = docx.Document(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text)
    return [Document(page_content=text, metadata={"source": path})]

# ---------------- App state ----------------
if "mm_records" not in st.session_state:
    st.session_state.mm_records = []
    st.session_state.vectors = []
    st.session_state.mm_index = None

splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=80)

def add_text_chunks(clip, docs: List[Document], src_meta: Dict[str, Any]):
    for d in splitter.split_documents(docs):
        st.session_state.mm_records.append({
            "id": uuid.uuid4().hex,
            "modality": "text",
            "text": d.page_content,
            "source": {**src_meta, **(d.metadata or {})},
        })
        st.session_state.vectors.append(embed_text(clip, d.page_content))

def add_image_item(clip, b: bytes, src_meta: Dict[str, Any]):
    img = Image.open(io.BytesIO(b)).convert("RGB")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tf:
        img.save(tf.name); img_path = tf.name
    st.session_state.mm_records.append({
        "id": uuid.uuid4().hex, "modality": "image",
        "image_path": img_path, "source": src_meta
    })
    st.session_state.vectors.append(embed_image(clip, img))

def add_audio_segments(clip, audio_path: str, result: Dict[str, Any], src_meta: Dict[str, Any]):
    for seg in result.get("segments", []):
        text = seg.get("text", "").strip()
        if not text:
            continue
        start = float(seg.get("start", 0.0))
        end = float(seg.get("end", 0.0))
        st.session_state.mm_records.append({
            "id": uuid.uuid4().hex, "modality": "audio_text",
            "text": text, "audio_path": audio_path,
            "start": start, "end": end,
            "source": {**src_meta, "segment": [start, end]},
        })
        st.session_state.vectors.append(embed_text(clip, text))

def rebuild_index():
    if not st.session_state.vectors:
        st.session_state.mm_index = None
        return
    mat = np.vstack(st.session_state.vectors).astype("float32")
    faiss.normalize_L2(mat)
    idx = faiss.IndexFlatIP(mat.shape[1]); idx.add(mat)
    st.session_state.mm_index = idx

def search_by_text(clip, q: str, k: int = 12):
    if st.session_state.mm_index is None: return []
    qv = embed_text(clip, q)
    D, I = st.session_state.mm_index.search(qv[None, :], k)
    return [(st.session_state.mm_records[i], float(D[0][j]), i) for j,i in enumerate(I[0]) if i != -1]

def search_by_image(clip, b: bytes, k: int = 12):
    if st.session_state.mm_index is None: return []
    img = Image.open(io.BytesIO(b)).convert("RGB")
    qv = embed_image(clip, img)
    D, I = st.session_state.mm_index.search(qv[None, :], k)
    return [(st.session_state.mm_records[i], float(D[0][j]), i) for j,i in enumerate(I[0]) if i != -1]

def grounded_answer(llm, question: str, hits, max_tokens=512):
    ctx, srcs = [], []
    for i, (rec, score, _) in enumerate(hits, 1):
        sid = f"[{i}]"
        snippet = rec.get("text", "")[:1000] if rec["modality"] != "image" else "<image>"
        ctx.append(f"{sid} modality={rec['modality']} score={score:.3f} SOURCE={json.dumps(rec.get('source', {}))}\n{snippet}")
        srcs.append({"num": i, **rec})
    sys = ("You are a grounded assistant. Use ONLY the provided CONTEXT. "
           "Cite facts with [n]. If unknown, say you don't know.")
    prompt = f"{sys}\n\nCONTEXT:\n{chr(10).join(ctx)}\n\nQUESTION: {question}\n\nAnswer clearly with [n]. Limit to {max_tokens} tokens.\n"
    return Ollama(model="llama3").invoke(prompt), srcs

# ---------------- Sidebar ----------------
clip = load_clip_model()
with st.sidebar:
    st.header("📥 Ingest Sources")
    d_docs  = st.file_uploader("Upload documents (PDF/TXT/DOCX)", type=["pdf","txt","docx"], accept_multiple_files=True)
    d_imgs  = st.file_uploader("Upload images (PNG/JPG/WebP)", type=["png","jpg","jpeg","webp"], accept_multiple_files=True)
    d_audio = st.file_uploader("Upload audio (WAV/MP3/M4A/WEBM/OGG)", type=["wav","mp3","m4a","webm","ogg"], accept_multiple_files=True)

    st.markdown("---")
    url_input = st.text_input("🌐 Fetch URL (optional)")
    yt_url    = st.text_input("📺 YouTube URL (optional)")

    st.markdown("---")
    max_hits = st.slider("Top-K retrieval", 3, 20, 8)
    max_answer_tokens = st.slider("Answer length (tokens)", 128, 1500, 500, 32)

    if st.button("🧱 Rebuild Index"):
        rebuild_index(); st.success("Index rebuilt.")

# ---- Documents ----
if d_docs:
    for f in d_docs:
        ext = f.name.split(".")[-1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tf:
            tf.write(f.read()); path = tf.name
        try:
            if ext == "pdf":
                docs = PyPDFLoader(path).load()
            elif ext == "txt":
                docs = TextLoader(path).load()
            elif ext == "docx":
                docs = Docx2txtLoader(path).load() if HAS_DOCX2TXT else load_docx_via_python_docx(path)
            else:
                continue
            add_text_chunks(clip, docs, {"file": f.name, "type": ext})
        except Exception as e:
            st.error(f"Failed to load {f.name}: {e}")

# ---- Images ----
if d_imgs:
    for f in d_imgs:
        try:
            add_image_item(clip, f.read(), {"file": f.name, "type": "image"})
        except Exception as e:
            st.error(f"Failed to add image {f.name}: {e}")

# ---- Audio uploads (PyAV decode) ----
if d_audio:
    if not HAS_WHISPER:
        st.error("Whisper not installed. Run: pip install whisper openai-whisper")
    else:
        wh = load_whisper_model()
        for f in d_audio:
            with tempfile.NamedTemporaryFile(delete=False, suffix="."+f.name.split(".")[-1]) as tf:
                tf.write(f.read()); in_path = tf.name
            try:
                audio = decode_audio_to_16k_mono(in_path)
                if audio.size == 0:
                    raise RuntimeError("File contains no decodable audio.")
                result = wh.transcribe(audio, fp16=False)
                wav_full = os.path.join(tempfile.gettempdir(), uuid.uuid4().hex + ".wav")
                write_wav(wav_full, audio, TARGET_SR)
                add_audio_segments(clip, wav_full, result, {"file": f.name, "type": "audio"})
            except Exception as e:
                st.error(f"Failed to process audio {f.name}: {e}")

# ---- URL ----
if url_input:
    try:
        res = requests.get(url_input, timeout=8)
        soup = BeautifulSoup(res.text, "html.parser")
        content = soup.get_text("\n")
        add_text_chunks(clip, [Document(page_content=content, metadata={"url": url_input})], {"url": url_input, "type": "html"})
    except Exception as e:
        st.error(f"Failed to fetch URL: {e}")

# ---- YouTube (robust: format fallback + validation + timeout + max duration) ----
def download_youtube_best_audio(url: str, tmpd: str, max_secs: int, progress_cb=None) -> str:
    """
    Download a playable audio file for the given YouTube URL.
    Tries several format selectors (prefers m4a). Validates acodec and file existence.
    Returns local file path. Raises on failure.
    """
    if not HAS_YTDLP:
        raise RuntimeError("yt-dlp not installed")

    from yt_dlp.utils import DownloadError, match_filter_func

    # Prefer m4a (AAC) first – PyAV handles it well; then generic bestaudio
    format_candidates = [
        "bestaudio[ext=m4a]/bestaudio/best",
        "bestaudio[acodec!=none]/bestaudio/best",
        "bestaudio/best",
    ]

    def _hook(d):
        if progress_cb:
            progress_cb(d)

    for fmt in format_candidates:
        outtmpl = os.path.join(tmpd, "yt_audio.%(ext)s")
        ydl_opts = {
            "format": fmt,
            "outtmpl": outtmpl,
            "quiet": True,
            "noplaylist": True,
            "socket_timeout": 20,
            "retries": 3,
            "progress_hooks": [_hook],
            "match_filter": match_filter_func(f"duration <= {max_secs}"),
            "concurrent_fragment_downloads": 1,
            "nocheckcertificate": True,
        }
        try:
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=True)
                # quick sanity: ensure audio is present
                if info.get("acodec") == "none" or info.get("audio_ext") == "none":
                    raise RuntimeError(f"Selected format has no audio (fmt={fmt})")
                path = ydl.prepare_filename(info)
            if not os.path.exists(path):
                raise RuntimeError("Downloaded file not found on disk.")
            return path
        except DownloadError as e:
            # try next selector
            continue
        except Exception:
            # try next selector
            continue

    raise RuntimeError("Could not download a valid audio format (video may be restricted or too long).")

if yt_url:
    if not HAS_YTDLP:
        st.error("yt-dlp not installed. Run: pip install yt-dlp")
    elif not HAS_WHISPER:
        st.error("Whisper not installed. Run: pip install whisper openai-whisper")
    else:
        MAX_SECS = 10 * 60  # 10 minutes cap
        prog_txt = st.empty()

        def hook(d):
            if d.get("status") == "downloading":
                prog_txt.info(f"Downloading… {d.get('_percent_str','').strip()} | ETA: {d.get('eta','?')}s")
            elif d.get("status") == "finished":
                prog_txt.info("Download finished. Decoding & transcribing…")

        try:
            with st.status("Downloading and transcribing YouTube audio…", expanded=False):
                with tempfile.TemporaryDirectory() as tmpd:
                    src_path = download_youtube_best_audio(yt_url, tmpd, MAX_SECS, progress_cb=hook)
                    audio = decode_audio_to_16k_mono(src_path, max_seconds=MAX_SECS)
                if audio.size == 0:
                    raise RuntimeError("Downloaded file had no decodable audio track.")

                wh = load_whisper_model()
                result = wh.transcribe(audio, fp16=False)
                wav_full = os.path.join(tempfile.gettempdir(), uuid.uuid4().hex + ".wav")
                write_wav(wav_full, audio, TARGET_SR)
                add_audio_segments(clip, wav_full, result, {"url": yt_url, "type": "youtube"})
                st.success("YouTube audio processed.")
        except Exception as e:
            st.error(f"YouTube error: {e}")

# Build index if needed
if st.session_state.mm_index is None and st.session_state.vectors:
    rebuild_index()

# ---------------- Query UI ----------------
left, right = st.columns([2,1])
with left:
    st.subheader("🔎 Unified Query Interface")
    mode = st.radio("Query type", ["Text","Image","Audio"], horizontal=True)

    query_bytes, user_query = None, None
    if mode == "Text":
        user_query = st.text_input("Type your question (e.g., ‘Show the report mentioning international development in 2024’ or ‘find the screenshot at 14:32’)")
    elif mode == "Image":
        qi = st.file_uploader("Drop a query image (PNG/JPG/WebP)", type=["png","jpg","jpeg","webp"], accept_multiple_files=False, key="qimg")
        if qi:
            query_bytes = qi.read()
            st.image(query_bytes, caption="Query image", use_column_width=True)
    else:
        qa = st.file_uploader("Drop an audio query (WAV/MP3/M4A/WEBM/OGG)", type=["wav","mp3","m4a","webm","ogg"], accept_multiple_files=False, key="qaud")
        if qa and HAS_WHISPER:
            with tempfile.NamedTemporaryFile(delete=False, suffix="."+qa.name.split(".")[-1]) as tf:
                tf.write(qa.read()); qpath = tf.name
            st.audio(qpath)
            try:
                q_audio = decode_audio_to_16k_mono(qpath)
                if q_audio.size == 0:
                    raise RuntimeError("Query audio has no decodable samples.")
                wh = load_whisper_model()
                out = wh.transcribe(q_audio, fp16=False)
                user_query = out.get("text","").strip()
                if user_query: st.info(f"Audio transcript query: {user_query}")
            except Exception as e:
                st.error(f"Unable to transcribe query audio: {e}")

    do = st.button("Search & Answer")

with right:
    st.subheader("📦 Corpus Summary")
    st.write(f"**Indexed items:** {len(st.session_state.mm_records)}")
    if st.session_state.mm_records:
        counts = {"text":0,"image":0,"audio_text":0}
        for r in st.session_state.mm_records:
            counts[r["modality"]] = counts.get(r["modality"],0)+1
        st.json(counts)

# ---------------- Retrieval + Grounded answer ----------------
def show_hits(hits):
    st.markdown("### 📑 Results (by modality)")
    bucket = {"text":[], "image":[], "audio_text":[]}
    for rec, score, idx in hits:
        bucket[rec["modality"]].append((rec, score, idx))

    if bucket["text"]:
        with st.expander(f"📝 Text matches ({len(bucket['text'])})", expanded=True):
            for i,(rec,score,_) in enumerate(bucket["text"],1):
                st.write(f"**{i}. score={score:.3f}** · **source:**"); st.code(json.dumps(rec.get("source", {})))
                st.write("> " + rec.get("text","")[:500].replace("\n"," ") + "…")

    if bucket["image"]:
        with st.expander(f"🖼️ Image matches ({len(bucket['image'])})", expanded=True):
            for i,(rec,score,_) in enumerate(bucket["image"],1):
                st.write(f"**{i}. score={score:.3f}** · **source:**"); st.code(json.dumps(rec.get("source", {})))
                st.image(rec.get("image_path"), use_column_width=True)

    if bucket["audio_text"]:
        with st.expander(f"🔊 Audio transcript matches ({len(bucket['audio_text'])})", expanded=True):
            for i,(rec,score,_) in enumerate(bucket["audio_text"],1):
                st.write(f"**{i}. score={score:.3f}** · **segment:** {rec.get('start',0):.2f}–{rec.get('end',0):.2f}s · **source:**")
                st.code(json.dumps(rec.get("source", {})))
                try:
                    full, sr = sf.read(rec["audio_path"], dtype="float32")
                    s = max(int(rec["start"]*TARGET_SR)-200, 0)
                    e = min(int(rec["end"]*TARGET_SR)+200, len(full))
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tf:
                        sf.write(tf.name, full[s:e], TARGET_SR, subtype="PCM_16")
                        st.audio(tf.name)
                except Exception as e:
                    st.caption(f"(Could not render snippet: {e})")

if do:
    if st.session_state.mm_index is None or not st.session_state.vectors:
        st.warning("No index yet. Please upload some data and click 'Rebuild Index'.")
    else:
        if mode == "Text" and user_query:
            hits = search_by_text(clip, user_query, k=max_hits)
        elif mode == "Image" and query_bytes:
            hits = search_by_image(clip, query_bytes, k=max_hits)
        elif mode == "Audio" and user_query:
            hits = search_by_text(clip, user_query, k=max_hits)
        else:
            hits = []

        if not hits:
            st.info("No results.")
        else:
            show_hits(hits)
            ans, cites = grounded_answer(Ollama(model="llama3"), user_query if user_query else "Describe matches",
                                         hits[:max_hits], max_tokens=max_answer_tokens)
            st.markdown("### 🧠 Grounded Answer"); st.write(ans)

            st.markdown("### 🔗 Citations & Source Navigation")
            for s in cites:
                st.markdown(f"**[{s['num']}]** · **{s['modality']}**"); st.code(json.dumps(s.get("source", {})))
                if s["modality"] == "text":
                    st.code(s.get("text","")[:1200])
                elif s["modality"] == "image":
                    st.image(s.get("image_path"), use_column_width=True)
                elif s["modality"] == "audio_text":
                    st.caption(f"Segment: {s.get('start',0):.2f}–{s.get('end',0):.2f}s")
                    try:
                        full, sr = sf.read(s["audio_path"], dtype="float32")
                        a = max(int(s["start"]*TARGET_SR)-200,0)
                        b = min(int(s["end"]*TARGET_SR)+200, len(full))
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tf:
                            sf.write(tf.name, full[a:b], TARGET_SR, subtype="PCM_16")
                            st.audio(tf.name)
                    except Exception as e:
                        st.caption(f"(Could not render snippet: {e})")

